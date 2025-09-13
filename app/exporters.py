"""DOCX export utilities."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def export_docx(sections: list[dict], citations: list[dict], path: Path) -> None:
    """Write a simple SR&ED report to ``path``.

    Parameters
    ----------
    sections:
        List of dicts with ``{"section": str, "text": str}``.
    citations:
        List of dicts with ``{"path", "start", "end", "text"}`` keys.
    path:
        Destination file path.
    """

    doc = Document()
    doc.add_heading("SR&ED Draft", level=0)

    for section in sections:
        doc.add_heading(section["section"], level=1)
        doc.add_paragraph(section["text"]) 

    doc.add_page_break()
    doc.add_heading("Evidence Appendix", level=1)
    for cit in citations:
        doc.add_paragraph(
            f"{cit['path']}:{cit['start']}-{cit['end']} — {cit['text'].strip()}"
        )

    doc.save(path)

